
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

document = Document()

url = 'https://www.vedantu.com/formula/calculus-formulas'
response = requests.get(url)
#print(response.text)

soup = BeautifulSoup(response.content,'html.parser')
#print(response.text)

data=soup.find_all('span',style ='font-size: 13pt;  color: rgb(0, 0, 0); background-color: transparent; font-variant-numeric: normal; font-variant-east-asian: normal; vertical-align: baseline; white-space: pre-wrap;')
#########store data in word file

document.add_heading('Math Note', 0)

for i in data:
    i=i.text.strip()
    document.add_paragraph(i)

#####save file
document.save('demo.docx')



